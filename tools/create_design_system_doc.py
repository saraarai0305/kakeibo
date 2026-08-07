from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"D:\仕事用\5_その他\mainichi\design\LIFE_NOTE_デザインシステムと実装ロードマップ_v1.docx")

COLORS = {
    "ink": "132238",
    "muted": "5F7189",
    "line": "D9E2EC",
    "surface": "F6F8FB",
    "blue": "4B9FE8",
    "green": "34B98B",
    "coral": "F27D78",
    "violet": "9577DD",
    "yellow": "F1BE48",
    "danger": "D95055",
    "white": "FFFFFF",
}


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_keep(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Yu Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(p, text, size=10.5, bold=False, color=None):
    run = p.add_run(text)
    set_font(run, size, bold, color)
    return run


def add_para(doc, text="", size=10.5, bold=False, color=None, space_after=5, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.42
    if text:
        add_text(p, text, size, bold, color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    add_text(p, text, 18 if level == 1 else 13.5, True, COLORS["ink"])
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    add_text(p, text, 10.5, False, COLORS["ink"])
    return p


def add_callout(doc, label, text, color):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.7)
    table.columns[1].width = Cm(13.8)
    left, right = table.rows[0].cells
    set_cell_shading(left, color)
    set_cell_shading(right, "F7FAFC")
    for cell in (left, right):
        set_cell_margins(cell, 120, 160, 120, 160)
        set_cell_border(cell, top={"val": "single", "sz": 6, "color": color}, bottom={"val": "single", "sz": 6, "color": color}, left={"val": "single", "sz": 6, "color": color}, right={"val": "single", "sz": 6, "color": color})
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, label, 10, True, COLORS["white"])
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, text, 10.5, False, COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Cm(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, COLORS["ink"])
        set_cell_margins(cell, 110, 120, 110, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, h, 9.5, True, COLORS["white"])
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_shading(cell, "FFFFFF" if r_i % 2 == 0 else "F7FAFC")
            set_cell_margins(cell, 105, 120, 105, 120)
            set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": COLORS["line"]})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            add_text(p, str(value), 9.6, False, COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_checklist(doc, items):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(0.7)
    table.columns[1].width = Cm(15.8)
    for label, text in items:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_shading(cell, "FFFFFF")
            set_cell_margins(cell, 90, 80, 90, 80)
            set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": COLORS["line"]})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text(p, label, 11, True, COLORS["blue"])
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, text, 10.2, False, COLORS["ink"])


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    add_text(p, "LIFE NOTE  |  デザインシステムと実装ロードマップ  |  v1.0", 8.5, False, COLORS["muted"])


def style_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(1.65)
    sec.bottom_margin = Cm(1.55)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)
    for section in doc.sections:
        add_footer(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)


def build():
    doc = Document()
    style_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(45)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "LIFE NOTE", 13, True, COLORS["blue"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(10)
    add_text(p, "デザインシステムと\n実装ロードマップ", 27, True, COLORS["ink"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    add_text(p, "情報を最初に見せすぎず、目的から必要な一画面へ進める生活管理アプリ", 11, False, COLORS["muted"])

    cover = doc.add_table(rows=1, cols=3)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.autofit = False
    for i, (label, val, color) in enumerate([
        ("設計の順番", "基準 → 3画面 → 横展開", COLORS["blue"]),
        ("対象", "ホーム / お金 / 一日の流れ", COLORS["violet"]),
        ("判断軸", "少ない・迷わない・続く", COLORS["green"]),
    ]):
        cell = cover.rows[0].cells[i]
        set_cell_shading(cell, "F7FAFC")
        set_cell_margins(cell, 170, 140, 170, 140)
        set_cell_border(cell, top={"val": "single", "sz": 8, "color": color}, bottom={"val": "single", "sz": 4, "color": COLORS["line"]}, left={"val": "single", "sz": 4, "color": COLORS["line"]}, right={"val": "single", "sz": 4, "color": COLORS["line"]})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        add_text(p, label, 9, True, COLORS["muted"])
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text(p, val, 10, True, COLORS["ink"])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_callout(doc, "決定", "以後の実装は「デザインシステムに合うか」を先に確認し、合わない既存UIは残さず置き換える。画面単位の継ぎ足し修正はしない。", COLORS["ink"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    add_text(p, "2026年8月7日  |  v1.0", 9.5, False, COLORS["muted"])

    doc.add_page_break()

    # 1. rationale
    add_heading(doc, "1. この計画で解決すること")
    add_para(doc, "現在の課題は、機能が不足していることではなく、情報・操作・装飾が同じ強さで並び、利用者が「今すること」を選びにくいことです。以後は、画面ごとの目的を一つに絞り、情報は必要になった段階で開く構造に変えます。", 10.8, False, COLORS["ink"], 8)
    add_table(doc, ["現状の問題", "置き換える原則"], [
        ("1画面に複数の主役がある", "一画面につき、主目的と主アクションは各1つにする"),
        ("同じ操作が旧UI・新UIに重複する", "新しい導線を入れる時点で古い導線を削除する"),
        ("枠とカードが多く、読む順番が散る", "基本は余白と区切り線。面で強調するのは主役だけ"),
        ("文字・アイコン・色の規則がページごとに揺れる", "数値トークンと単一アイコンセットを全画面で使う"),
        ("アニメーションが装飾で、意味を伝えない", "状態変化・保存・展開にだけ、短く一貫した動きを使う"),
    ], [7.6, 8.9])

    add_heading(doc, "2. 情報アーキテクチャ（入口 → 目的 → 独立画面）")
    add_para(doc, "ホームは情報を読ませるページではなく、生活の入口です。数値・進捗・説明を置かず、次の三択だけを大きく見せます。", 10.5, False, COLORS["ink"], 5)
    add_table(doc, ["入口", "次の選択", "到達する独立画面", "その画面の主役"], [
        ("記録する", "支出・収入 / こころとからだ", "支出・収入を記録 / 体調を記録", "保存する"),
        ("今日を整える", "一日の流れ / テーマ", "今日の時間割 / テーマ設定", "今日を確認・編集する"),
        ("見える化する", "お金 / こころとからだ", "お金の分析 / 体調の分析", "傾向を読む"),
    ], [2.4, 3.7, 6.0, 4.4])
    add_callout(doc, "禁止", "目的を選んだ後、別目的の情報を同じ画面に混在させない。たとえば一日の流れには支出分析や習慣の集計を常時置かない。", COLORS["coral"])

    # 3 system
    add_heading(doc, "3. 固定するデザインシステム")
    add_para(doc, "これは見た目の好みではなく、迷わず実装・レビューするための共通言語です。例外が必要な場合は、先にこの資料を更新してから実装します。", 10.5, False, COLORS["ink"], 6)
    add_heading(doc, "3-1. 余白・文字・操作の基準", 2)
    add_table(doc, ["要素", "固定ルール", "目的"], [
        ("余白", "4px単位。主に 8 / 16 / 24 / 32px を使用", "視線のまとまりを作る"),
        ("文字", "12 / 14 / 16 / 20 / 28 / 36px の6段階だけ", "階層を明確にする"),
        ("本文", "16px未満にしない。行高は1.4〜1.6", "iPhoneで無理なく読める"),
        ("主ボタン", "高さ52px以上、文字16px以上", "迷わず押せる"),
        ("補助ボタン", "高さ44px以上、文字14px以上", "タップ領域を保証する"),
        ("区切り", "情報群の区切りは余白16〜24pxまたは1px罫線", "カードの過剰使用を防ぐ"),
    ], [2.3, 8.0, 4.7])
    add_heading(doc, "3-2. 色の意味", 2)
    add_table(doc, ["用途", "色", "使う場所", "使わない場所"], [
        ("お金（中立）", "青 #4B9FE8", "金額入力・お金の遷移", "支出と収入の区別"),
        ("支出・マイナス", "赤 #D95055", "支出タブ・負の金額・注意", "装飾目的の背景"),
        ("収入・プラス", "緑 #34B98B", "収入タブ・正の金額・達成", "常時強調した本文"),
        ("からだ", "緑 #34B98B", "身体評価・歩数・睡眠", "こころの評価"),
        ("こころ", "コーラル #F27D78", "気分・心の評価", "金銭の状態"),
        ("こよみ", "紫 #9577DD", "予定・時間割・カレンダー", "警告"),
        ("期限・注意", "黄 #F1BE48", "支払期日・未処理の注意", "主要ボタン"),
    ], [2.4, 2.6, 5.2, 4.8])
    add_heading(doc, "3-3. アイコンと文字の細部", 2)
    add_bullet(doc, "アイコンは単一の線画SVGセットに統一する。絵文字・3Dアイコン・別テイストの記号は使用しない。")
    add_bullet(doc, "標準サイズは24px、座布団は48px。アイコンはCSS上の中心だけでなく、実機スクリーンショットで視覚中心を確認する。")
    add_bullet(doc, "数字は可能な限り tabular-nums を使用し、金額は「¥4,000」、歩数は「2,619歩」のように単位まで一行で完結させる。")
    add_bullet(doc, "日本語本文には任意の広い字間を使わない。見出しは字間0〜0.02em、英字のLIFE NOTEだけ0.12em程度に限定する。")

    # 4 principles
    add_heading(doc, "4. デザインの四原則を実装判断へ落とす")
    add_table(doc, ["原則", "実装上の約束", "レビュー時の質問"], [
        ("近接", "同じ目的の情報は16px以内に集め、目的が変わる場所は24px以上離す", "これは同じ操作のための情報か？"),
        ("整列", "見出し・本文・金額・アイコンの開始線を揃え、異なるグリッドを同じ画面に混在させない", "視線は左上から自然に流れるか？"),
        ("反復", "ボタン、セグメント、リスト行、グラフ凡例の形と挙動を全画面で再利用する", "同じ役割の部品が同じ見え方か？"),
        ("対比", "強い面・大きな数字・鮮やかな色は主役1つに限定する", "一目で最初に見る場所が決まるか？"),
    ], [1.8, 8.7, 4.5])
    add_callout(doc, "判定", "どの要素も「なぜここで目立つのか」を説明できない場合は削除または補助階層へ下げる。", COLORS["violet"])

    # 5 central screens
    doc.add_page_break()
    add_heading(doc, "5. 最初に作り直す3つの中心画面")
    add_para(doc, "この3画面を先に完成させ、部品・余白・アニメーション・戻る導線を実機で検証します。残りの画面はここで確定した部品だけで作ります。", 10.5, False, COLORS["ink"], 8)
    add_heading(doc, "5-1. ホーム：目的を選ぶハブ", 2)
    add_table(doc, ["目的", "構成", "残すもの", "削除するもの"], [
        ("次にすることを3秒で決める", "LIFE NOTE / 質問 / 三択", "記録する・今日を整える・見える化する", "数値、進捗、注意書き、設定の重複導線"),
    ], [3.6, 4.0, 4.2, 3.2])
    add_bullet(doc, "三択は縦に並べるが、巨大なカードにはしない。1行の大型リストボタンとして、アイコン・名称・矢印を同じ基準線に置く。")
    add_bullet(doc, "上部は「戻る」を置かない。ホームだけが起点であり、設定は右上の44pxボタンから開く。")
    add_heading(doc, "5-2. 支出・収入を記録：金額を残す画面", 2)
    add_table(doc, ["目的", "視線の順番", "主アクション", "補助"], [
        ("支出または収入を迷わず保存する", "今日使えるお金 → 今残っているお金 → 支出/収入 → 金額 → 方法 → カテゴリー", "支出を記録 / 収入を記録", "方法・カテゴリーを追加する"),
    ], [3.6, 6.8, 3.1, 2.0])
    add_bullet(doc, "今日使えるお金は最上段の主役。今残っているお金は次の行で全額表示し、表示／非表示だけを切り替える。")
    add_bullet(doc, "支出は赤、収入は緑。タブ、保存ボタン、金額の符号に同じ意味を反復する。")
    add_bullet(doc, "入力は縦一列。「金額 → 方法 → カテゴリー → 保存」の順に固定し、横並び入力を使わない。")
    add_heading(doc, "5-3. 一日の流れ：時間を読む画面", 2)
    add_table(doc, ["目的", "構成", "常時表示", "開いて表示"], [
        ("現在時刻と今日の予定を時間順に確認・編集する", "日付とテーマ → 時間割 → 3操作", "赤い現在時刻線・時間・予定", "予定追加・チェックリスト・週/月カレンダー"),
    ], [3.5, 4.2, 4.1, 3.7])
    add_bullet(doc, "現在時刻は予定の上に重ねず、時間目盛りから右へ伸びる1pxの赤線と右端の時刻ラベルだけで示す。")
    add_bullet(doc, "予定を足す／チェックリストを開く／週・月を見るは時間割の外側に置く。時間軸の中に絶対配置しない。")
    add_bullet(doc, "週・月カレンダーは初期状態で閉じ、押したときだけボトムシートまたは下部展開で開く。")

    # 6 component/motion
    add_heading(doc, "6. 部品とアニメーションの規格")
    add_table(doc, ["部品", "規格", "状態"], [
        ("ページヘッダー", "戻る / 中央タイトル / 右操作。高さ56px。開始線とタップ領域を固定", "戻るは前画面へ、ホームへは別の明示操作"),
        ("大型リストボタン", "高さ72px。24pxアイコン、16px名称、右矢印", "押下時: 120msで背景のみ変化"),
        ("主ボタン", "高さ52px。16px太字。画面内に原則1つ", "保存中→完了を短く表示"),
        ("セグメント", "高さ44px。選択状態は文字＋背景＋色で示す", "タップ直後に内容を切替"),
        ("入力行", "ラベル → 52px入力欄を縦に配置", "エラーは項目直下に表示"),
        ("グラフ", "見出し・期間・凡例・グラフ・補助値の順", "凡例タップで表示ON/OFF"),
    ], [3.0, 8.1, 4.4])
    add_heading(doc, "6-1. モーション仕様", 2)
    add_table(doc, ["場面", "時間", "動き", "意図"], [
        ("ボタン押下", "120ms", "明度と1px縮小のみ", "タップを受け取ったことを伝える"),
        ("画面遷移", "200ms", "右から入る/左へ戻る。ease-out", "現在地と戻る方向を伝える"),
        ("展開・折りたたみ", "180ms", "高さと不透明度を同時に変化", "情報がどこから現れたかを伝える"),
        ("保存完了", "240ms", "ボタン文言を完了に置換し短く色変更", "保存された安心感を与える"),
        ("数値・グラフ", "240ms", "表示時のみ上昇・描画", "変化を読む補助。自動ループ禁止"),
    ], [3.0, 2.0, 5.3, 5.2])
    add_bullet(doc, "すべての動きは prefers-reduced-motion を尊重し、必要時は即時表示へ切り替える。")
    add_bullet(doc, "常時動く装飾、意味のないバウンス、画面を待たせるローディングは禁止する。処理待ちが発生した時だけ、小さなローディング表示を使う。")

    # 7 roadmap
    doc.add_page_break()
    add_heading(doc, "7. 実装ロードマップと優先順位")
    add_para(doc, "後から機能を足すためではなく、戻せない複雑さを増やさないための順番です。各段階の完了条件を満たすまで次へ進みません。", 10.5, False, COLORS["ink"], 7)
    add_table(doc, ["優先", "作業", "成果物", "完了条件"], [
        ("P0", "現状棚卸し・旧UI削除", "画面一覧、旧要素の削除リスト", "同じ役割の操作が1つだけ。未使用DOM/CSS/JSを削除"),
        ("P1", "デザイントークン・共通部品", "CSS変数、文字・余白・ボタン・アイコン規格", "3画面で同じ部品が同じ見た目と挙動"),
        ("P2", "ホーム再設計", "3つの大型リストボタンを持つハブ", "3秒以内に目的を選べる。余白が過剰でない"),
        ("P3", "支出・収入記録の再設計", "縦順入力・金額表示・方法/カテゴリー追加", "金額が省略されず、支出/収入の色が一貫"),
        ("P4", "一日の流れ再設計", "時間割、赤い現在時刻線、折りたたみ導線", "時間とボタンが重ならず、旧操作が残らない"),
        ("P5", "横展開", "体調・分析・設定・支払い予定", "P1部品のみで構成。目的外の情報を常時出さない"),
        ("P6", "実機QAと公開", "iPhone画面確認、デプロイ確認", "スクロール量・中心揃え・保存・戻るを全画面確認"),
    ], [1.2, 3.4, 4.8, 6.1])
    add_callout(doc, "順番", "P0〜P4が終わるまでは、分析の細かな追加機能や新しいカードを増やさない。先に骨格と反復を完成させる。", COLORS["yellow"])

    # 8 QA
    add_heading(doc, "8. 各画面で必ず行うレビュー")
    add_checklist(doc, [
        ("□", "画面名と主目的を、1文で説明できる"),
        ("□", "最初に見るべき要素が1つだけ決まっている"),
        ("□", "主アクションは1つで、52px以上の高さがある"),
        ("□", "補助操作は44px以上で、主役を邪魔していない"),
        ("□", "アイコン、文字、矢印が同一の基準線・視覚中心に揃っている"),
        ("□", "金額・日付・単位が途中で切れず、一行で意味を読める"),
        ("□", "支出=赤、収入=緑、その他の意味色が混ざっていない"),
        ("□", "旧UI、重複ボタン、説明文、未使用要素が残っていない"),
        ("□", "戻る操作が直前の画面へ戻り、ホームへ戻る操作と混同しない"),
        ("□", "iPhoneの実機幅でスクリーンショットを撮り、切れ・重なり・中心ずれがない"),
    ])
    add_heading(doc, "9. 実装開始時の運用ルール")
    add_bullet(doc, "画面を変更する前に、この資料のP0〜P6のどこに当たるかを明記する。")
    add_bullet(doc, "変更後はPCだけで判断せず、iPhone幅のスクリーンショットで文字・アイコン・時間軸を確認する。")
    add_bullet(doc, "公開前に、古いDOM、古いボタン文言、旧イベントハンドラ、未使用スタイルを検索し、残存がないことを確認する。")
    add_bullet(doc, "データ構造と同期方式は維持し、UIの置き換えで既存のiPhoneデータを消さない。")
    add_bullet(doc, "この資料にない例外が必要なら、先に理由・範囲・代替案を追記して合意を取る。")
    add_callout(doc, "次の作業", "P0として、現行UIを画面・DOM・操作単位で棚卸しし、残すもの／削除するもの／P1で置き換えるものを一覧化する。", COLORS["blue"])

    # Metadata
    doc.core_properties.title = "LIFE NOTE デザインシステムと実装ロードマップ v1"
    doc.core_properties.subject = "生活管理アプリの情報設計、UI規格、実装優先順位"
    doc.core_properties.author = "LIFE NOTE"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
